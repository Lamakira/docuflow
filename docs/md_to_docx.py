"""Minimal Markdown to DOCX for DocuFlow docs (headings, lists, tables, bold)."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_formatted_runs(paragraph, text):
    """Split **bold** and add runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for p in parts:
        if p.startswith("**") and p.endswith("**"):
            run = paragraph.add_run(p[2:-2])
            run.bold = True
        elif p:
            paragraph.add_run(p)


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and "|" in line[1:]


def parse_table_row(line: str):
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def main():
    md_path = Path(__file__).with_name("DocuFlow-User-Test-Guide-EN.md")
    out_path = Path(__file__).with_name("DocuFlow-User-Test-Guide-EN.docx")
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        # Table block
        if is_table_row(line):
            table_buffer = [line]
            i += 1
            while i < len(lines) and is_table_row(lines[i]):
                table_buffer.append(lines[i])
                i += 1
            rows = [parse_table_row(r) for r in table_buffer if not re.match(r"^\|[\s\-:|]+\|$", r.strip())]
            sep_re = re.compile(r"^\|[\s\-:|]+\|$")
            rows = [parse_table_row(r) for r in table_buffer if not sep_re.match(r.strip())]
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_formatted_runs(p, cell_text)
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # List items
        m = re.match(r"^(\s*)- \[([ xX])\]\s+(.*)$", line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ " if m.group(2) == " " else "☑ ")
            add_formatted_runs(p, text)
            i += 1
            continue

        m = re.match(r"^(\s*)-\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, m.group(2))
            i += 1
            continue

        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, m.group(2))
            i += 1
            continue

        # Empty
        if not line.strip():
            i += 1
            continue

        # Italic-only line
        if line.startswith("*") and line.endswith("*") and line.count("*") == 2:
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, line.strip())
        i += 1

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
